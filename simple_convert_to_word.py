#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
简化的Markdown转Word脚本
"""

import sys
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def simple_markdown_to_word(markdown_file, word_file):
    """简化的Markdown转Word函数"""
    
    print(f"正在转换: {markdown_file} → {word_file}")
    
    # 读取Markdown文件
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 创建Word文档
    doc = Document()
    
    # 设置文档属性
    doc.core_properties.title = "Mac mini + 飞书 + Deepseek Chat 部署指导书"
    doc.core_properties.author = "Jarvis AI助手"
    doc.core_properties.subject = "技术部署指南"
    doc.core_properties.keywords = "Mac mini, 飞书, Deepseek, AI, 部署, 对接"
    doc.core_properties.comments = "详细的技术部署指导文档"
    
    # 添加封面
    add_simple_cover(doc)
    
    # 解析内容
    lines = content.split('\n')
    in_code_block = False
    code_content = []
    
    for line in lines:
        line = line.rstrip()
        
        if line.startswith('```'):
            if in_code_block:
                # 代码块结束
                if code_content:
                    add_code_block(doc, '\n'.join(code_content))
                    code_content = []
                in_code_block = False
            else:
                # 代码块开始
                in_code_block = True
            continue
        
        if in_code_block:
            # 收集代码内容
            code_content.append(line)
            continue
            
        if not line:
            # 空行
            doc.add_paragraph()
            continue
            
        # 处理标题
        if line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_heading(title, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif line.startswith('## '):
            title = line[3:].strip()
            doc.add_heading(title, level=1)
            
        elif line.startswith('### '):
            title = line[4:].strip()
            doc.add_heading(title, level=2)
            
        elif line.startswith('#### '):
            title = line[5:].strip()
            doc.add_heading(title, level=3)
            
        elif line.startswith('- ') or line.startswith('* '):
            # 列表项
            list_text = line[2:].strip()
            doc.add_paragraph(list_text, style='List Bullet')
            
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
            # 编号列表
            list_text = line[3:].strip()
            doc.add_paragraph(list_text, style='List Number')
            
        elif line.startswith('    ') or line.startswith('\t'):
            # 代码或缩进内容
            code_text = line.strip()
            if code_text:
                p = doc.add_paragraph(code_text)
                p.style = 'Normal'
                p.paragraph_format.left_indent = Inches(0.5)
                
        else:
            # 普通段落
            if line.strip():
                p = doc.add_paragraph(line)
                p.style = 'Normal'
    
    # 保存文档
    doc.save(word_file)
    print(f"✅ Word文档已创建: {word_file}")
    
    return word_file

def add_simple_cover(doc):
    """添加简单封面"""
    
    # 添加大标题
    title = doc.add_heading('Mac mini + 飞书 + Deepseek Chat', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 空行
    
    subtitle = doc.add_heading('大模型部署对接详细指导书', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加多个空行
    for _ in range(15):
        doc.add_paragraph()
    
    # 添加版本信息
    version = doc.add_paragraph(f"版本: v1.0")
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 空行
    
    # 添加日期
    date = doc.add_paragraph(f"创建日期: {datetime.now().strftime('%Y年%m月%d日')}")
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 空行
    
    # 添加作者
    author = doc.add_paragraph("作者: Jarvis AI助手")
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 分页
    doc.add_page_break()
    
    # 添加目录标题
    doc.add_heading('目录', level=1)
    
    # 简化的目录
    toc_items = [
        "第一部分：系统架构概述",
        "第二部分：Mac mini 环境准备", 
        "第三部分：Deepseek Chat 大模型部署",
        "第四部分：飞书开放平台集成",
        "第五部分：系统集成与对接",
        "第六部分：部署与运维",
        "第七部分：最佳实践与案例",
        "附录"
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.style = 'Normal'
    
    doc.add_page_break()

def add_code_block(doc, code_text):
    """添加代码块"""
    if not code_text.strip():
        return
    
    # 创建代码段落
    p = doc.add_paragraph()
    p.style = 'Normal'
    
    # 添加代码文本
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(39, 55, 70)
    
    # 设置缩进
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def main():
    """主函数"""
    
    # 输入输出文件
    markdown_file = "/Users/jiyingguo/.openclaw/workspace/complete_guide.md"
    word_file = "/Users/jiyingguo/.openclaw/workspace/Mac_mini_飞书_Deepseek_部署指导书.docx"
    
    try:
        # 检查输入文件
        if not os.path.exists(markdown_file):
            print(f"错误: Markdown文件不存在: {markdown_file}")
            sys.exit(1)
        
        # 转换文档
        output_file = simple_markdown_to_word(markdown_file, word_file)
        
        print("\n" + "="*60)
        print("文档转换完成！")
        print("="*60)
        print(f"输出文件: {output_file}")
        print(f"文件大小: {os.path.getsize(output_file) / 1024:.1f} KB")
        
        return output_file
        
    except Exception as e:
        print(f"转换失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()