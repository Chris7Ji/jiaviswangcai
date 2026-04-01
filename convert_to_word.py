#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将Markdown转换为Word文档
"""

import sys
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def markdown_to_word(markdown_file, word_file):
    """将Markdown文件转换为Word文档"""
    
    print(f"正在转换: {markdown_file} → {word_file}")
    
    # 读取Markdown文件
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 创建Word文档
    doc = Document()
    
    # 设置文档属性
    doc.core_properties.title = "Mac mini + 飞书 + Deepseek Chat 部署指导书"
    doc.core_properties.author = "Jarvis AI助手"
    doc.core_properties.subject = "技术部署指南"
    doc.core_properties.keywords = "Mac mini, 飞书, Deepseek, AI, 部署, 对接"
    doc.core_properties.comments = "详细的技术部署指导文档"
    
    # 创建自定义样式
    styles = doc.styles
    
    # 标题1样式
    title1 = styles.add_style('CustomTitle1', WD_STYLE_TYPE.PARAGRAPH)
    title1.font.name = 'Microsoft YaHei'
    title1.font.size = Pt(24)
    title1.font.bold = True
    title1.font.color.rgb = RGBColor(0, 0, 0)
    
    # 标题2样式
    title2 = styles.add_style('CustomTitle2', WD_STYLE_TYPE.PARAGRAPH)
    title2.font.name = 'Microsoft YaHei'
    title2.font.size = Pt(18)
    title2.font.bold = True
    title2.font.color.rgb = RGBColor(44, 62, 80)
    
    # 标题3样式
    title3 = styles.add_style('CustomTitle3', WD_STYLE_TYPE.PARAGRAPH)
    title3.font.name = 'Microsoft YaHei'
    title3.font.size = Pt(14)
    title3.font.bold = True
    title3.font.color.rgb = RGBColor(52, 73, 94)
    
    # 正文样式
    normal = styles['Normal']
    normal.font.name = 'Microsoft YaHei'
    normal.font.size = Pt(11)
    
    # 代码样式
    code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(10)
    code_style.font.color.rgb = RGBColor(39, 55, 70)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)
    
    # 添加封面页
    add_cover_page(doc)
    
    # 添加目录
    add_table_of_contents(doc)
    
    # 解析Markdown内容
    lines = content.split('\n')
    current_heading_level = 0
    
    for line in lines:
        line = line.rstrip()
        
        if not line:
            # 空行
            doc.add_paragraph()
            continue
            
        # 检查标题
        if line.startswith('# '):
            # 标题1
            title = line[2:].strip()
            p = doc.add_paragraph(title, style='CustomTitle1')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            current_heading_level = 1
            
        elif line.startswith('## '):
            # 标题2
            title = line[3:].strip()
            doc.add_paragraph(title, style='CustomTitle2')
            current_heading_level = 2
            
        elif line.startswith('### '):
            # 标题3
            title = line[4:].strip()
            doc.add_paragraph(title, style='CustomTitle3')
            current_heading_level = 3
            
        elif line.startswith('```'):
            # 代码块开始
            continue
            
        elif line.startswith('    ') or line.startswith('\t'):
            # 代码或缩进内容
            code_text = line.strip()
            if code_text:
                p = doc.add_paragraph(code_text, style='Code')
                
        elif line.startswith('- ') or line.startswith('* '):
            # 列表项
            list_text = line[2:].strip()
            p = doc.add_paragraph(list_text, style='Normal')
            p.paragraph_format.left_indent = Inches(0.25)
            
        elif line.startswith('|'):
            # 表格行
            if '---' in line:
                # 表格分隔线，跳过
                continue
            else:
                # 表格内容
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                if cells:
                    # 简化处理：作为普通段落
                    p = doc.add_paragraph(' | '.join(cells), style='Normal')
                    
        else:
            # 普通段落
            if line.strip():
                p = doc.add_paragraph(line, style='Normal')
    
    # 添加页脚
    add_footer(doc)
    
    # 保存文档
    doc.save(word_file)
    print(f"✅ Word文档已创建: {word_file}")
    
    return word_file

def add_cover_page(doc):
    """添加封面页"""
    
    # 添加大标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = title.add_run("Mac mini + 飞书 + Deepseek Chat")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(28)
    run.font.bold = True
    
    doc.add_paragraph()  # 空行
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = subtitle.add_run("大模型部署对接详细指导书")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(22)
    run.font.bold = True
    
    # 添加多个空行
    for _ in range(10):
        doc.add_paragraph()
    
    # 添加版本信息
    version_info = doc.add_paragraph()
    version_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = version_info.add_run("版本: v1.0")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(14)
    
    doc.add_paragraph()  # 空行
    
    # 添加日期
    date_info = doc.add_paragraph()
    date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = date_info.add_run(f"创建日期: {datetime.now().strftime('%Y年%m月%d日')}")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(14)
    
    doc.add_paragraph()  # 空行
    
    # 添加作者
    author_info = doc.add_paragraph()
    author_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = author_info.add_run("作者: Jarvis AI助手")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(14)
    
    # 分页
    doc.add_page_break()

def add_table_of_contents(doc):
    """添加目录"""
    
    toc_title = doc.add_paragraph("目录")
    toc_title.style = 'CustomTitle2'
    
    # 简化的目录内容
    toc_items = [
        "第一部分：系统架构概述",
        "第二部分：Mac mini 环境准备", 
        "第三部分：Deepseek Chat 大模型部署",
        "第四部分：飞书开放平台集成",
        "第五部分：系统集成与对接",
        "第六部分：部署与运维",
        "第七部分：最佳实践与案例",
        "附录"
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item, style='Normal')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 分页
    doc.add_page_break()

def add_footer(doc):
    """添加页脚"""
    
    # 获取所有节
    for section in doc.sections:
        footer = section.footer
        
        # 添加页脚内容
        paragraph = footer.paragraphs[0]
        paragraph.text = f"Mac mini + 飞书 + Deepseek Chat 部署指导书 - 第 "
        
        # 添加页码
        run = paragraph.add_run()
        run.add_field('PAGE', '1')
        
        run = paragraph.add_run(" 页")
        
        # 添加总页数
        run = paragraph.add_run(" / ")
        run.add_field('NUMPAGES', '1')
        run = paragraph.add_run(" 页")
        
        # 添加日期
        run = paragraph.add_run(f" - {datetime.now().strftime('%Y年%m月%d日')}")
        
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def main():
    """主函数"""
    
    # 输入输出文件
    markdown_file = "/Users/jiyingguo/.openclaw/workspace/complete_guide.md"
    word_file = "/Users/jiyingguo/.openclaw/workspace/Mac_mini_飞书_Deepseek_部署指导书.docx"
    
    try:
        # 检查输入文件
        if not os.path.exists(markdown_file):
            print(f"错误: Markdown文件不存在: {markdown_file}")
            sys.exit(1)
        
        # 转换文档
        output_file = markdown_to_word(markdown_file, word_file)
        
        print("\n" + "="*60)
        print("文档转换完成！")
        print("="*60)
        print(f"输出文件: {output_file}")
        print(f"文件大小: {os.path.getsize(output_file) / 1024:.1f} KB")
        
        return output_file
        
    except Exception as e:
        print(f"转换失败: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()